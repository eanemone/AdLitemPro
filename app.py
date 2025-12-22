import streamlit as st
import os
import logging
import urllib.parse
import pickle
import tiktoken
import re
import time
from io import BytesIO

# --- NEW: WORD DOC GENERATION ---
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- MODERN PARTNER PACKAGES ---
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings, ChatOpenAI

# --- STANDARD LANGCHAIN IMPORTS ---
from langchain.retrievers import ParentDocumentRetriever, EnsembleRetriever
from langchain.storage import LocalFileStore
from langchain.storage._lc_store import create_kv_docstore
from langchain_community.retrievers import BM25Retriever

# --- CORE BUILDING BLOCKS ---
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_text_splitters import RecursiveCharacterTextSplitter

# --- CONFIGURATION ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ST_DB_PATH = os.path.join(BASE_DIR, "legal_db_vectors")
DOC_STORE_PATH = os.path.join(BASE_DIR, "legal_docstore_fs")
BM25_PATH = os.path.join(BASE_DIR, "bm25_retriever.pkl")
COLLECTION_NAME = "legal_cases_eyecite"

PREFERRED_MODEL = "gpt-4o" 
TEMPERATURE = 0.4

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("AdLitemPro")

# --- UI SETUP ---
st.set_page_config(page_title="AdLitem Pro", layout="wide", page_icon="⚖️")

# --- CUSTOM CSS (NUCLEAR BLUE THEME) ---
st.markdown("""
<style>
    .stApp { max-width: 1100px; margin: 0 auto; }
    
    /* BRANDING & HEADERS */
    .main-header { font-family: 'Helvetica Neue', sans-serif; font-size: 2.8rem; color: #FFFFFF; font-weight: 800; text-align: center; margin-bottom: 0.2rem; }
    .subtitle { font-size: 0.95rem; color: #94A3B8; text-align: center; margin-bottom: 2rem; font-weight: 400; letter-spacing: 0.05em; }
    
    /* --- AGGRESSIVE INPUT STYLING (THE RED KILLER) --- */
    /* Target the Chat Input specifically */
    div[data-testid="stChatInput"] textarea:focus {
        border-color: #38BDF8 !important;
        box-shadow: 0 0 0 1px #38BDF8 !important;
    }
    /* Target standard text inputs */
    div[data-baseweb="input"]:focus-within {
        border-color: #38BDF8 !important;
        box-shadow: 0 0 0 1px #38BDF8 !important;
    }
    /* Target ALL textareas and inputs just to be safe */
    textarea:focus, input:focus {
        border-color: #38BDF8 !important;
        box-shadow: 0 0 0 1px #38BDF8 !important;
    }
    textarea, input {
        color: #FFFFFF !important;
    }
    /* Button Styling */
    .stButton button {
        border-color: #38BDF8 !important;
        color: #38BDF8 !important;
    }
    .stButton button:hover {
        border-color: #0EA5E9 !important;
        color: #0EA5E9 !important;
    }

    /* MEMO STYLES */
    .memo-container { background: #FFFFFF; border-radius: 8px; border: 1px solid #E2E8F0; overflow: hidden; margin-bottom: 1.5rem; box-shadow: 0 2px 10px rgba(0,0,0,0.05); }
    .section-header { background: linear-gradient(135deg, #0369A1 0%, #0284C7 100%); color: #FFFFFF; padding: 12px 24px; font-family: 'Helvetica Neue', sans-serif; font-size: 1.1rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.08em; border-bottom: 3px solid #38BDF8; }
    .question-presented { background-color: #F8FAFC; color: #0F172A; padding: 24px; border-bottom: 1px solid #E2E8F0; font-family: 'Georgia', serif; font-size: 1.05rem; line-height: 1.6; }
    .brief-answer { background-color: #F8FAFC; color: #0F172A; padding: 24px; border-bottom: 1px solid #E2E8F0; font-family: 'Georgia', serif; font-size: 1.05rem; line-height: 1.6; }
    .discussion-box { background-color: #FFFFFF; color: #1E293B; padding: 32px; font-family: 'Georgia', serif; font-size: 1.1rem; line-height: 1.8; }
    .memo-header { color: #0369A1; font-weight: 800; font-size: 1.4rem; margin-top: 1.5rem; margin-bottom: 0.8rem; font-family: 'Helvetica Neue', sans-serif; text-transform: uppercase; letter-spacing: 0.03em; }
    .subsection-header { color: #0369A1; font-weight: 700; font-size: 1.2rem; margin-top: 1.8rem; margin-bottom: 0.6rem; font-family: 'Helvetica Neue', sans-serif; text-transform: uppercase; letter-spacing: 0.05em; border-bottom: 2px solid #38BDF8; padding-bottom: 0.3rem; }
    .inline-citation { color: #0284c7; font-weight: bold; font-size: 0.9em; }

    /* AUTHORITY LIST */
    .auth-item { border-left: 4px solid #38BDF8; background: #1E293B; padding: 15px; margin-bottom: 12px; border-radius: 0 4px 4px 0; }
    .auth-label { font-size: 0.7rem; font-weight: 800; color: #38BDF8; text-transform: uppercase; margin-bottom: 5px; letter-spacing: 0.05em; }
    .auth-title { font-weight: 700; color: #F1F5F9; font-size: 1rem; margin-bottom: 2px; }
    .auth-cite { color: #94A3B8; font-size: 0.85rem; font-style: italic; margin-bottom: 8px; }
    .auth-snippet { color: #CBD5E1; font-size: 0.85rem; line-height: 1.4; border-top: 1px solid #334155; padding-top: 8px; margin-top: 8px; }
    .scholar-link-inline { color: #38BDF8; text-decoration: none; font-size: 0.75rem; font-weight: 600; margin-left: 10px; }
</style>
""", unsafe_allow_html=True)

# --- AUTHENTICATION GATE ---
def check_password():
    def password_entered():
        if "username" in st.session_state and "password" in st.session_state:
            user = st.session_state["username"]
            pwd = st.session_state["password"]
            if (user in st.secrets["passwords"] and pwd == st.secrets["passwords"][user]):
                st.session_state["password_correct"] = True
                del st.session_state["password"]
                del st.session_state["username"]
            else:
                st.session_state["password_correct"] = False

    credit_html = """<div style="text-align: center; margin-bottom: 15px;"><span style="font-family: 'Helvetica Neue', sans-serif; font-size: 0.75rem; color: #64748B; letter-spacing: 0.1em; text-transform: uppercase;">Created by <span style="color: #E2E8F0; font-weight: 600;">Ernest Anemone, Esq.</span></span></div>"""

    if "password_correct" not in st.session_state:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown('<div style="margin-top: 80px;"></div>', unsafe_allow_html=True)
            st.markdown('<div class="main-header">AdLitem<span style="color:#38BDF8">Pro</span></div>', unsafe_allow_html=True)
            st.markdown(credit_html, unsafe_allow_html=True)
            st.markdown('<div class="subtitle">AUTHORIZED PERSONNEL ONLY</div>', unsafe_allow_html=True)
            st.text_input("Username", key="username")
            st.text_input("Password", type="password", on_change=password_entered, key="password")
        return False
    elif not st.session_state["password_correct"]:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown('<div style="margin-top: 80px;"></div>', unsafe_allow_html=True)
            st.markdown('<div class="main-header">AdLitem<span style="color:#38BDF8">Pro</span></div>', unsafe_allow_html=True)
            st.markdown(credit_html, unsafe_allow_html=True)
            st.markdown('<div class="subtitle">AUTHORIZED PERSONNEL ONLY</div>', unsafe_allow_html=True)
            st.text_input("Username", key="username")
            st.text_input("Password", type="password", on_change=password_entered, key="password")
            st.error("😕 Access Denied")
        return False
    return True

if not check_password():
    st.stop()

# --- UTILITIES ---
def clean_plain_text(text: str) -> str:
    if not text: return ""
    clean = re.compile('<.*?>')
    return re.sub(clean, '', str(text)).strip()

def clean_llm_output(text: str) -> str:
    text = re.sub(r'^```html\s*', '', text)
    text = re.sub(r'\s*```$', '', text)
    return text.strip()

# --- HEADER CLEANER & CITATION ENFORCER ---
def enforce_citations(text: str) -> str:
    """Finds missed citations and wraps them in HTML span tags."""
    statute_pattern = r'(?<!class="inline-citation">)(N\.J\.A\.C\.|N\.J\.S\.A\.|N\.J\.|N\.J\. Super\.)\s*(\d+[:\-]\d+[\d\-\.\w]*)'
    text = re.sub(statute_pattern, r'<span class="inline-citation">\1 \2</span>', text, flags=re.IGNORECASE)
    
    # Update to match CPP format
    policy_pattern = r'(?<!class="inline-citation">)(CPP[-\s][IVX\d\-\w]+)'
    text = re.sub(policy_pattern, r'<span class="inline-citation">\1</span>', text, flags=re.IGNORECASE)
    
    # Also catch CIC Manual references
    cic_pattern = r'(?<!class="inline-citation">)(CIC Manual\s*§?\s*\d+\.\d+)'
    text = re.sub(cic_pattern, r'<span class="inline-citation">\1</span>', text, flags=re.IGNORECASE)
    
    return text

def strip_redundant_headers(text: str) -> str:
    """Removes AI-generated headers (e.g. '**Brief Answer**') that duplicate our UI headers."""
    lines = text.split('\n')
    cleaned_lines = []
    
    # Matches: **Brief Answer**, ## Discussion, Brief Answer, etc.
    redundant_pattern = re.compile(r'^[\*\#\s]*(Brief Answer|Discussion)[\*\#\s]*$', re.IGNORECASE)
    
    for line in lines:
        if redundant_pattern.match(line.strip()):
            continue # Skip this line
        cleaned_lines.append(line)
        
    return "\n".join(cleaned_lines)

# --- WORD DOC GENERATOR (WITH MARKDOWN PARSING) ---
def create_docx(content: str) -> BytesIO:
    doc = Document()
    
    # Title
    title = doc.add_heading('Legal Research Memo', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated by AdLitem Pro | {time.strftime('%B %d, %Y')}")
    doc.add_paragraph("__________________________________________________________________")

    # Clean Content first (Remove duplicates)
    content = strip_redundant_headers(content)

    # Split content
    parts = content.split("===SECTION_BREAK===")
    full_text_lines = []
    
    if len(parts) > 1:
        full_text_lines.append("BRIEF ANSWER")
        full_text_lines.extend(parts[0].split('\n'))
        full_text_lines.append("") 
        full_text_lines.append("DISCUSSION")
        full_text_lines.extend(parts[1].split('\n'))
    else:
        full_text_lines = content.split('\n')

    # Regex Config
    header_re = re.compile(r'<div class="memo-header">(.*?)</div>', re.IGNORECASE)
    citation_re = re.compile(r'<span class="inline-citation">(.*?)</span>', re.IGNORECASE)
    bold_re = re.compile(r'\*\*(.*?)\*\*') # Finds **Bold Text**
    
    for line in full_text_lines:
        line = line.strip()
        if not line: continue
            
        # Headers
        if header_re.search(line):
            doc.add_heading(header_re.search(line).group(1), level=1)
            continue
        if line in ["BRIEF ANSWER", "DISCUSSION"]:
            doc.add_heading(line, level=1)
            continue

        # Clean HTML Citations for Word
        clean_line = citation_re.sub(r'\1', line) 
        clean_line = clean_plain_text(clean_line)
        
        # Paragraph + Markdown Bolding
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Parse "**" bolding
        segments = bold_re.split(clean_line)
        for i, segment in enumerate(segments):
            if not segment: continue
            run = p.add_run(segment)
            if i % 2 == 1: # Odd segments are inside **markers**
                run.bold = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def render_memo_ui(content: str, key_idx: int):
    # 1. Strip redundant headers so UI is clean
    content = strip_redundant_headers(content)
    
    # 2. Enforce Blue Citations
    content = enforce_citations(content)
    
    # 3. Render HTML with proper section headers
    if "===SECTION_BREAK===" in content:
        parts = content.split("===SECTION_BREAK===")
        question = parts[0].strip() if len(parts) > 0 else ""
        brief = parts[1].strip() if len(parts) > 1 else ""
        disc = parts[2].strip() if len(parts) > 2 else ""
        
        st.markdown(f'''
            <div class="memo-container">
                <div class="section-header">Question Presented</div>
                <div class="question-presented">{question}</div>
                <div class="section-header">Brief Answer</div>
                <div class="brief-answer">{brief}</div>
                <div class="section-header">Discussion</div>
                <div class="discussion-box">{disc}</div>
            </div>
        ''', unsafe_allow_html=True)
    else:
        st.markdown(f'<div class="memo-container"><div class="discussion-box">{content}</div></div>', unsafe_allow_html=True)

    # 4. Download Button
    doc_file = create_docx(content)
    st.download_button(
        label="📄 Download Memo as Word Doc",
        data=doc_file,
        file_name=f"Legal_Memo_{key_idx}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"dl_btn_{key_idx}"
    )

def get_badge_label(metadata):
    dtype = metadata.get("type", "case")
    if dtype == "policy": return "DCF POLICY"
    if dtype == "statute": return "STATUTE"
    if dtype == "manual": return "CIC MANUAL"
    return "PUBLISHED" if metadata.get("is_published") else "UNPUBLISHED"

# --- IMPROVED CITATION FORMATTER ---
def format_citation(meta: dict) -> str:
    """
    Extract proper citation from metadata based on type.
    Now works with actual DB structure: display_name, source, type.
    """
    doc_type = meta.get("type", "case")
    
    # Try bluebook first (for cases)
    cite_str = clean_plain_text(meta.get("bluebook", ""))
    
    # If bluebook is empty, use display_name or source
    if not cite_str:
        cite_str = meta.get("display_name", "") or meta.get("source", "")
    
    # Type-specific formatting
    if doc_type == "manual":
        # CIC_Manual_1108 -> CIC Manual § 11.08
        if "CIC" in cite_str or "cic" in cite_str.lower():
            match = re.search(r'(\d{2})(\d{2})', cite_str)
            if match:
                return f"CIC Manual § {match.group(1)}.{match.group(2)}"
            return "NJ DCF CIC Manual"
    
    elif doc_type == "policy":
        # CPP-IV-E-1-1100 -> CP&P-IV-E-1-1100
        if "CPP" in cite_str or "cpp" in cite_str.lower():
            # Clean up the format
            cite_str = re.sub(r'\.pdf$', '', cite_str, flags=re.IGNORECASE)
            cite_str = cite_str.replace("CPP", "CP&P")
            return cite_str
    
    elif doc_type == "statute":
        # Extract N.J.A.C. or N.J.S.A. format
        if "njac" in cite_str.lower() or "n.j.a.c" in cite_str.lower():
            match = re.search(r'(\d+:\d+[\d\-\.]*)', cite_str)
            if match:
                return f"N.J.A.C. {match.group(1)}"
        elif "njsa" in cite_str.lower() or "n.j.s.a" in cite_str.lower():
            match = re.search(r'(\d+[A-Z]?:\d+[\d\-\.]*)', cite_str)
            if match:
                return f"N.J.S.A. {match.group(1)}"
    
    # Clean up file extensions if still present
    cite_str = re.sub(r'\.(pdf|txt)$', '', cite_str, flags=re.IGNORECASE)
    
    return cite_str if cite_str else "Authority"

# --- NEW: HISTORY CONTEXTUALIZATION ---
def rewrite_query(original_input, chat_history):
    if not chat_history: return original_input
    llm = ChatOpenAI(model=PREFERRED_MODEL, temperature=0.3)
    prompt = f"Chat History:\n{chat_history}\n\nFollow-Up Input: {original_input}\n\nStandalone Question:"
    return llm.invoke(prompt).content

# --- DATABASE LOADING ---
@st.cache_resource
def get_retriever():
    if not os.path.exists(ST_DB_PATH): 
        return None
    
    vectorstore = Chroma(
        collection_name=COLLECTION_NAME, 
        embedding_function=OpenAIEmbeddings(), 
        persist_directory=ST_DB_PATH
    )
    
    fs = LocalFileStore(DOC_STORE_PATH)
    store = create_kv_docstore(fs)
    
    pdr = ParentDocumentRetriever(
        vectorstore=vectorstore, 
        docstore=store, 
        child_splitter=RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100),
        search_kwargs={"k": 15}
    )
    
    if os.path.exists(BM25_PATH):
        try:
            with open(BM25_PATH, "rb") as f:
                bm25 = pickle.load(f)
                return EnsembleRetriever(retrievers=[pdr, bm25], weights=[0.7, 0.3])
        except Exception:
            return pdr
    return pdr

# --- MAIN APP UI ---
st.markdown('<div class="main-header">AdLitem<span style="color:#38BDF8">Pro</span></div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">NEW JERSEY CHILD WELFARE LAW RESEARCH ENGINE</div>', unsafe_allow_html=True)

if "messages" not in st.session_state: st.session_state.messages = []
if "last_sources" not in st.session_state: st.session_state.last_sources = []

# --- LANDING PAGE ---
if not st.session_state.messages:
    st.markdown('<div style="height: 4vh;"></div>', unsafe_allow_html=True)
    st.markdown("""<div style="text-align: center; margin-bottom: 40px;"><span style="font-family: 'Helvetica Neue', sans-serif; font-size: 0.85rem; color: #64748B; letter-spacing: 0.1em; text-transform: uppercase;">Created by <span style="color: #E2E8F0; font-weight: 600;">Ernest Anemone, Esq.</span></span></div>""", unsafe_allow_html=True)
    st.markdown("""<div style="text-align: center; color: #94A3B8; margin-bottom: 30px; font-size: 0.9rem;">Select a sample fact pattern to test the research engine:</div>""", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    def set_prompt(text): st.session_state.messages.append({"role": "user", "content": text})
    
    with col1:
        if st.button("🌿 Cannabis & 'Imminent Harm'", use_container_width=True):
            set_prompt("Analyze the impact of the NJ CREAMMA Act on Title 9 litigation. Specifically: Does a positive newborn toxicology screen for THC, absent evidence of actual parenting impairment, legally sustain a finding of abuse or neglect? Cite recent Appellate Division precedents.")
            st.rerun()
    with col2:
        if st.button("🚧 'Safety Plans' & Due Process", use_container_width=True):
            set_prompt("Analyze the legality of DCPP 'Safety Protection Plans' that require a parent to leave the home under threat of removal. Does this constitute a 'constructive removal' requiring a Dodd hearing?")
            st.rerun()
    with col3:
        if st.button("👨‍👩‍👧 KLG vs. Adoption Preference", use_container_width=True):
            set_prompt("If a resource parent unequivocally prefers adoption, can the court still grant Kinship Legal Guardianship (KLG) to avoid Termination of Parental Rights (TPR)? Analyze the 'clear and convincing' evidence standard under the KLG Act amendments.")
            st.rerun()

# Render History
for idx, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg["role"]):
        if msg["role"] == "assistant":
            render_memo_ui(msg["content"], idx)
        else:
            st.markdown(msg["content"])

# Input
if prompt := st.chat_input("Start legal research task..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    st.rerun()

# Execution
if st.session_state.messages and st.session_state.messages[-1]["role"] == "user":
    current_prompt = st.session_state.messages[-1]["content"]
    with st.chat_message("assistant"):
        retriever = get_retriever()
        if not retriever:
            st.error("Database not found. Please check repository file structure.")
        else:
            progress_bar = st.progress(0, text="Analyzing request history...")
            with st.status("Conducting Deep Research...", expanded=False) as status:
                chat_history_str = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.messages[:-1]])
                search_query = rewrite_query(current_prompt, chat_history_str)
                
                try:
                    progress_bar.progress(20, text="Querying database...")
                    docs = retriever.invoke(search_query)
                    supp_query = f"{search_query} DCF Policy CP&P N.J.A.C. N.J.S.A. Title 9 Title 30 CIC Manual"
                    supplemental = retriever.invoke(supp_query)[:15]
                    
                    unique_docs = []
                    seen = set()
                    for d in (docs + supplemental):
                        h = hash(d.page_content[:150])
                        if h not in seen:
                            seen.add(h)
                            unique_docs.append(d)
                    
                    unique_docs = unique_docs[:30]
                    context_blocks = []
                    st.session_state.last_sources = []
                    citation_map = {}

                    progress_bar.progress(50, text="Sanitizing authorities...")
                    for i, doc in enumerate(unique_docs):
                        meta = doc.metadata
                        
                        # Use the new citation formatter
                        cite_str = format_citation(meta)
                        
                        content = clean_plain_text(doc.page_content)
                        title = clean_plain_text(meta.get("display_name", "Authority"))
                        
                        context_blocks.append(f"SOURCE {i+1} [{meta.get('type','case').upper()} - {cite_str}]:\n{content}\n")
                        citation_map[i+1] = cite_str
                        
                        # Fixed Google Scholar link - use proper format
                        link = None
                        if meta.get("type", "case") == "case":
                            # Use the full citation for Scholar search
                            search_term = cite_str if cite_str else title
                            # Clean up the search term for URL encoding
                            search_term = re.sub(r'\s+', ' ', search_term).strip()
                            link = f"https://scholar.google.com/scholar?hl=en&as_sdt=4%2C31&q={urllib.parse.quote(search_term)}&oq="
                        
                        st.session_state.last_sources.append({
                            "label": get_badge_label(meta), 
                            "title": title, 
                            "cite": cite_str, 
                            "snippet": content[:350], 
                            "link": link
                        })
                        
                    status.update(label=f"Found {len(st.session_state.last_sources)} authorities.", state="complete")
                    progress_bar.progress(70, text="Drafting Research Memo...")

                    llm = ChatOpenAI(model=PREFERRED_MODEL, temperature=TEMPERATURE)
                    sys_prompt = """You are an Appellate Law Clerk. Write a formal Research Memo based ONLY on provided SOURCES.
                    
                    INTERPRETIVE APPROACH:
                    Before drafting, engage in hermeneutic analysis:
                    1. Consider the query holistically - understand the broader legal context and policy implications
                    2. Examine each source in light of the whole legal framework
                    3. Iteratively refine your understanding as you move between specific authorities and the general legal question
                    4. Recognize that each statute, case, and policy exists within a larger interpretive tradition
                    5. Identify tensions, ambiguities, and how different authorities inform each other
                    6. Let your understanding deepen through recursive engagement with the sources and question
                    
                    This interpretive work is internal - do not explicitly reference this process in your memo. Your final analysis should reflect this deeper understanding while maintaining traditional legal memo structure.
                    
                    STYLING RULES (CRITICAL):
                    1. You MUST cite your sources. Every claim must be followed by a citation from the provided CITATIONS list.
                    2. Format citations as: <span class="inline-citation">Actual Citation Text</span>. 
                    3. DO NOT use the word 'Cite' as a placeholder. Use the actual text from CITATIONS (e.g., 'CP&P-IV-E-1-1100' or 'CIC Manual § 11.08').
                    4. DO NOT split citations across lines.
                    5. Use the EXACT citation text provided in the CITATIONS map - do not abbreviate or modify it.
                    6. When citing policies, use the format 'CP&P-[section]' (e.g., CP&P-IV-E-1-1100).
                    7. When citing the CIC Manual, use 'CIC Manual § [section]' (e.g., CIC Manual § 11.08).
                    8. BLUEBOOK CITATION FORMAT: Always end sentences with a period BEFORE the citation. 
                       CORRECT: "The court held that removal was improper. <span class="inline-citation">N.J.S.A. 9:6-8.21</span>"
                       INCORRECT: "The court held that removal was improper <span class="inline-citation">N.J.S.A. 9:6-8.21</span>."
                    
                    MEMO STRUCTURE (MANDATORY):
                    You MUST structure the memo with THREE sections separated by '===SECTION_BREAK===':
                    
                    1. QUESTION PRESENTED
                       - State the legal question clearly and concisely
                       - Do NOT include a header (the UI handles it)
                    
                    ===SECTION_BREAK===
                    
                    2. BRIEF ANSWER
                       - Provide a direct yes/no or short answer
                       - Include key reasoning in 2-3 sentences
                       - Do NOT include a header (the UI handles it)
                    
                    ===SECTION_BREAK===
                    
                    3. DISCUSSION
                       - Do NOT include a "Discussion" header (the UI handles it)
                       - YOU MUST organize using these subsection headers:
                         <div class="subsection-header">Rule</div>
                         <div class="subsection-header">Analysis</div>
                         <div class="subsection-header">Conclusion</div>
                       - Rule: State the applicable legal standards and authorities
                       - Analysis: Apply the law to the facts with detailed citation
                       - Conclusion: Summarize the legal conclusion
                    
                    EXAMPLE OUTPUT STRUCTURE:
                    Whether [legal question]?
                    
                    ===SECTION_BREAK===
                    
                    Yes/No. [Brief reasoning with key cite].
                    
                    ===SECTION_BREAK===
                    
                    <div class="subsection-header">Rule</div>
                    [Legal standards and authorities]
                    
                    <div class="subsection-header">Analysis</div>
                    [Detailed application with citations]
                    
                    <div class="subsection-header">Conclusion</div>
                    [Final legal conclusion]
                    """
                    
                    chain = ChatPromptTemplate.from_messages([("system", sys_prompt), ("user", "CITATIONS: {citations}\n\nCONTEXT: {context}\n\nISSUE: {input}")]) | llm | StrOutputParser()
                    response = chain.invoke({"input": search_query, "context": "\n\n".join(context_blocks), "citations": citation_map})
                    
                    clean_resp = clean_llm_output(response)
                    progress_bar.progress(100, text="Memo Complete.")
                    time.sleep(0.5)
                    progress_bar.empty()
                    st.session_state.messages.append({"role": "assistant", "content": clean_resp})
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error: {e}")
                    progress_bar.empty()

# Footer
if st.session_state.last_sources:
    st.markdown("---")
    with st.expander("📚 View Authority Library", expanded=False):
        for src in st.session_state.last_sources:
            link = f'<a href="{src["link"]}" target="_blank" class="scholar-link-inline">View on Scholar ↗</a>' if src["link"] else ""
            st.markdown(f'<div class="auth-item"><div class="auth-label">{src["label"]}{link}</div><div class="auth-title">{src["title"]}</div><div class="auth-cite">{src["cite"]}</div><div class="auth-snippet">"{src["snippet"]}..."</div></div>', unsafe_allow_html=True)
    if st.button("🗑️ Clear Research Session", use_container_width=True):
        st.session_state.messages = []
        st.session_state.last_sources = []
        st.rerun()